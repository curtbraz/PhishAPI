from docx import Document
from docx.shared import Inches

import sys

UploadedDoc=sys.argv[1]

Path='/var/www/uploads/' + UploadedDoc

document = Document(Path)

document.add_picture('HTTP.png', width=0, height=0)

document.add_picture('UNC.png', width=0, height=0)

document.save('/var/www/uploads/Phishing.docx')
